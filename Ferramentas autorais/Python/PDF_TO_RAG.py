import argparse
import io
import json
import platform
import queue
import re
import shutil
import subprocess
import tempfile
import threading
import uuid
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

import fitz
import pytesseract
from PIL import Image
from docx import Document as DocxDocument

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}


def normalize_text(text):
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def smart_split(text, chunk_size=1200, overlap=200):
    if len(text) <= chunk_size:
        return [(0, len(text), text)]
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_size, text_len)
        if end < text_len:
            window_start = max(start, end - 250)
            window = text[window_start:end]
            split_candidates = [window.rfind('\n\n'), window.rfind('\n'), window.rfind('. '), window.rfind(' ')]
            best = max(split_candidates)
            if best > 100:
                end = window_start + best + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append((start, end, chunk))
        if end >= text_len:
            break
        start = max(end - overlap, start + 1)
    return chunks


def append_report(report, level, source, message, extra=None):
    item = {'level': level, 'source': str(source), 'message': str(message)}
    if extra is not None:
        item['extra'] = extra
    report.append(item)


def find_tesseract_cmd(custom_cmd=None):
    if custom_cmd:
        return custom_cmd
    found = shutil.which('tesseract')
    if found:
        return found
    if platform.system().lower().startswith('win'):
        common = Path(r'C:\Program Files\Tesseract-OCR\tesseract.exe')
        if common.exists():
            return str(common)
    return None


def configure_tesseract(custom_cmd=None):
    cmd = find_tesseract_cmd(custom_cmd)
    if cmd:
        pytesseract.pytesseract.tesseract_cmd = cmd
    return cmd


def pdf_page_to_ocr_text(page, ocr_lang='por+eng', dpi=250, tesseract_config=''):
    pix = page.get_pixmap(dpi=dpi, alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes('png')))
    text = pytesseract.image_to_string(img, lang=ocr_lang, config=tesseract_config)
    return normalize_text(text)


def extract_pdf_page_text(page):
    text = page.get_text('text') or ''
    return normalize_text(text)


def read_docx(path, report, original_source=None, original_name=None, original_type=None):
    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        append_report(report, 'error', path, f'Falha ao abrir DOCX: {e}')
        return []
    parts = []
    try:
        for p in doc.paragraphs:
            text = normalize_text(p.text)
            if text:
                parts.append(text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [normalize_text(cell.text) for cell in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    rows.append(' | '.join(cells))
            if rows:
                parts.append('\n'.join(rows))
    except Exception as e:
        append_report(report, 'error', path, f'Falha ao ler conteúdo do DOCX: {e}')
        return []
    text = normalize_text('\n\n'.join(parts))
    if not text:
        append_report(report, 'warning', path, 'DOCX sem texto extraível.')
        return []
    return [{
        'source': str(original_source or path),
        'file_name': original_name or path.name,
        'file_type': original_type or 'docx',
        'page': None,
        'text': text,
        'ocr_used': False,
        'ocr_reason': None,
    }]


def convert_doc_to_docx_windows(doc_path, output_dir):
    import pythoncom
    import win32com.client
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0
    output_path = Path(output_dir) / f'{Path(doc_path).stem}.docx'
    try:
        doc = word.Documents.Open(str(Path(doc_path).resolve()))
        doc.SaveAs(str(output_path.resolve()), FileFormat=16)
        doc.Close(False)
    finally:
        word.Quit()
        pythoncom.CoUninitialize()
    if not output_path.exists():
        raise RuntimeError('Conversão .doc -> .docx falhou no Word.')
    return output_path


def convert_doc_to_docx_unix(doc_path, output_dir):
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        raise RuntimeError('LibreOffice/Soffice não encontrado no PATH.')
    cmd = [
        soffice,
        '--headless',
        '--convert-to',
        'docx',
        '--outdir',
        str(output_dir),
        str(doc_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    output_path = Path(output_dir) / f'{Path(doc_path).stem}.docx'
    if result.returncode != 0 or not output_path.exists():
        raise RuntimeError(f'Falha na conversão com soffice. STDOUT: {result.stdout} STDERR: {result.stderr}')
    return output_path


def read_doc(path, report):
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            if platform.system().lower().startswith('win'):
                converted = convert_doc_to_docx_windows(path, tmpdir)
            else:
                converted = convert_doc_to_docx_unix(path, tmpdir)
        except Exception as e:
            append_report(report, 'error', path, f'Falha ao converter .doc para .docx: {e}')
            return []
        return read_docx(path=converted, report=report, original_source=path, original_name=path.name, original_type='doc')


class Processor:
    def __init__(self, ui_queue, settings):
        self.ui_queue = ui_queue
        self.settings = settings
        self.report = []
        self.documents = []
        self.chunks = []
        self.total_supported_files = 0
        self.processed_files = 0
        self.ocr_pages = 0
        self.stop_requested = False
        self.file_summaries = []

    def emit(self, kind, **data):
        self.ui_queue.put({'kind': kind, **data})

    def log(self, text):
        self.emit('log', text=text)

    def set_status(self, text):
        self.emit('status', text=text)

    def set_file(self, text):
        self.emit('file', text=text)

    def set_page(self, text):
        self.emit('page', text=text)

    def set_total_progress(self, current, total):
        self.emit('total_progress', current=current, total=total)

    def set_page_progress(self, current, total):
        self.emit('page_progress', current=current, total=total)

    def request_stop(self):
        self.stop_requested = True

    def gather_files(self, input_dir):
        input_path = Path(input_dir)
        files = []
        for path in input_path.rglob('*'):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                files.append(path)
        return sorted(files, key=lambda p: str(p).lower())

    def process(self):
        try:
            input_dir = self.settings['input_dir']
            output_json = self.settings['output_json']
            report_json = self.settings['report_json']
            chunk_size = self.settings['chunk_size']
            overlap = self.settings['overlap']
            ocr_lang = self.settings['ocr_lang']
            ocr_dpi = self.settings['ocr_dpi']
            ocr_min_chars = self.settings['ocr_min_chars']
            tesseract_config = self.settings['tesseract_config']
            tesseract_cmd = self.settings['tesseract_cmd']

            if not Path(input_dir).exists():
                raise FileNotFoundError(f'Pasta não encontrada: {input_dir}')

            cmd = configure_tesseract(tesseract_cmd)
            if cmd:
                self.log(f'Tesseract: {cmd}')
            else:
                append_report(self.report, 'warning', input_dir, 'Tesseract não encontrado. OCR não estará disponível.')
                self.log('Tesseract não encontrado. OCR ficará indisponível.')

            files = self.gather_files(input_dir)
            self.total_supported_files = len(files)

            if self.total_supported_files == 0:
                raise RuntimeError('Nenhum arquivo PDF, DOCX ou DOC encontrado.')

            self.set_total_progress(0, self.total_supported_files)
            self.set_page_progress(0, 1)
            self.set_status('Processando arquivos...')

            for index, path in enumerate(files, start=1):
                if self.stop_requested:
                    self.log('Processamento interrompido pelo usuário.')
                    break

                self.set_file(str(path))
                self.set_page('-')
                self.log(f'Lendo arquivo {index}/{self.total_supported_files}: {path.name}')

                ext = path.suffix.lower()
                file_entries_before = len(self.documents)
                file_ocr_before = self.ocr_pages
                file_summary = {
                    'source': str(path),
                    'file_name': path.name,
                    'file_type': ext.lstrip('.'),
                    'entries_added': 0,
                    'ocr_pages': 0,
                    'status': 'ok'
                }

                try:
                    if ext == '.pdf':
                        docs = self.read_pdf_with_progress(
                            path=path,
                            ocr_lang=ocr_lang,
                            ocr_dpi=ocr_dpi,
                            tesseract_config=tesseract_config,
                            ocr_min_chars=ocr_min_chars,
                            ocr_available=bool(cmd),
                        )
                        self.documents.extend(docs)
                    elif ext == '.docx':
                        docs = read_docx(path=path, report=self.report)
                        self.documents.extend(docs)
                        self.set_page_progress(1, 1)
                    elif ext == '.doc':
                        docs = read_doc(path=path, report=self.report)
                        self.documents.extend(docs)
                        self.set_page_progress(1, 1)
                except Exception as e:
                    append_report(self.report, 'error', path, f'Erro inesperado ao processar arquivo: {e}')
                    self.log(f'Erro em {path.name}: {e}')
                    file_summary['status'] = 'error'

                file_summary['entries_added'] = len(self.documents) - file_entries_before
                file_summary['ocr_pages'] = self.ocr_pages - file_ocr_before
                if file_summary['entries_added'] == 0 and file_summary['status'] == 'ok':
                    file_summary['status'] = 'empty'
                self.file_summaries.append(file_summary)

                self.processed_files = index
                self.set_total_progress(self.processed_files, self.total_supported_files)

            if self.documents:
                self.set_status('Gerando chunks...')
                self.log('Gerando chunks...')
                self.chunks = self.build_chunks(self.documents, chunk_size=chunk_size, overlap=overlap)
            else:
                self.chunks = []

            output_data = {
                'total_files_found': self.total_supported_files,
                'total_files_processed': self.processed_files,
                'total_text_entries': len(self.documents),
                'total_chunks': len(self.chunks),
                'ocr_pages': self.ocr_pages,
                'chunks': self.chunks,
            }

            with open(output_json, 'w', encoding='utf-8') as f:
                json.dump(output_data, f, ensure_ascii=False, indent=2)

            summary = {
                'output': output_json,
                'report': report_json,
                'total_files_found': self.total_supported_files,
                'total_files_processed': self.processed_files,
                'total_text_entries': len(self.documents),
                'total_chunks': len(self.chunks),
                'ocr_pages': self.ocr_pages,
                'files': self.file_summaries,
                'messages': self.report,
            }

            with open(report_json, 'w', encoding='utf-8') as f:
                json.dump(summary, f, ensure_ascii=False, indent=2)

            self.set_status('Concluído')
            self.set_file('-')
            self.set_page('-')
            self.set_page_progress(1, 1)

            self.emit(
                'done',
                output=output_json,
                report=report_json,
                total_files_found=self.total_supported_files,
                total_files_processed=self.processed_files,
                total_text_entries=len(self.documents),
                total_chunks=len(self.chunks),
                ocr_pages=self.ocr_pages,
            )
        except Exception as e:
            self.emit('error', text=str(e))

    def read_pdf_with_progress(self, path, ocr_lang, ocr_dpi, tesseract_config, ocr_min_chars, ocr_available):
        items = []
        doc = None
        try:
            doc = fitz.open(str(path))
        except Exception as e:
            append_report(self.report, 'error', path, f'Falha ao abrir PDF: {e}')
            self.log(f'Falha ao abrir PDF {path.name}: {e}')
            return items

        total_pages = len(doc)
        if total_pages == 0:
            append_report(self.report, 'warning', path, 'PDF sem páginas.')
            self.log(f'{path.name}: PDF sem páginas.')
            doc.close()
            return items

        self.set_page_progress(0, total_pages)
        try:
            for i, page in enumerate(doc, start=1):
                if self.stop_requested:
                    break
                self.set_page(f'{i}/{total_pages}')
                self.set_page_progress(i, total_pages)
                text = ''
                ocr_used = False
                ocr_reason = None
                try:
                    text = extract_pdf_page_text(page)
                except Exception as e:
                    ocr_reason = f'extract_exception: {e}'
                    append_report(self.report, 'warning', path, f'Falha na extração da página {i}, tentando OCR.', {'page': i, 'error': str(e)})
                    self.log(f'{path.name} | página {i}: extração falhou, tentando OCR.')
                if len(text.strip()) < ocr_min_chars:
                    if not ocr_reason:
                        ocr_reason = 'too_short_or_empty'
                    if ocr_available:
                        try:
                            text = pdf_page_to_ocr_text(page=page, ocr_lang=ocr_lang, dpi=ocr_dpi, tesseract_config=tesseract_config)
                            ocr_used = True
                            self.ocr_pages += 1
                            append_report(self.report, 'info', path, f'OCR usado na página {i}.', {'page': i, 'reason': ocr_reason})
                            self.log(f'{path.name} | página {i}: OCR usado.')
                        except Exception as e:
                            append_report(self.report, 'error', path, f'OCR falhou na página {i}.', {'page': i, 'error': str(e)})
                            self.log(f'{path.name} | página {i}: OCR falhou.')
                            text = ''
                    else:
                        append_report(self.report, 'warning', path, f'OCR necessário na página {i}, mas Tesseract não está disponível.', {'page': i, 'reason': ocr_reason})
                        self.log(f'{path.name} | página {i}: texto insuficiente e OCR indisponível.')
                if text:
                    items.append({
                        'source': str(path),
                        'file_name': path.name,
                        'file_type': 'pdf',
                        'page': i,
                        'text': text,
                        'ocr_used': ocr_used,
                        'ocr_reason': ocr_reason,
                    })
                else:
                    append_report(self.report, 'warning', path, f'Página {i} sem texto após tentativa de extração/OCR.', {'page': i})
                    self.log(f'{path.name} | página {i}: sem texto útil.')
        finally:
            doc.close()
        return items

    def build_chunks(self, documents, chunk_size=1200, overlap=200):
        chunks = []
        for doc in documents:
            pieces = smart_split(doc['text'], chunk_size=chunk_size, overlap=overlap)
            for idx, (char_start, char_end, piece) in enumerate(pieces):
                chunks.append({
                    'id': str(uuid.uuid4()),
                    'text': piece,
                    'source': doc['source'],
                    'file_name': doc['file_name'],
                    'file_type': doc['file_type'],
                    'page': doc['page'],
                    'chunk_index': idx,
                    'char_start': char_start,
                    'char_end': char_end,
                    'ocr_used': doc.get('ocr_used', False),
                    'ocr_reason': doc.get('ocr_reason'),
                })
        return chunks


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('PDF/DOCX/DOC → chunks.json')
        self.geometry('920x700')
        self.minsize(860, 640)
        self.ui_queue = queue.Queue()
        self.worker = None
        self.processor = None
        self.var_input_dir = tk.StringVar(value=str(Path.cwd()))
        self.var_output_json = tk.StringVar(value=str(Path.cwd() / 'chunks.json'))
        self.var_report_json = tk.StringVar(value=str(Path.cwd() / 'report.json'))
        self.var_chunk_size = tk.StringVar(value='1200')
        self.var_overlap = tk.StringVar(value='200')
        self.var_ocr_lang = tk.StringVar(value='por+eng')
        self.var_ocr_dpi = tk.StringVar(value='250')
        self.var_ocr_min_chars = tk.StringVar(value='30')
        self.var_tesseract_cmd = tk.StringVar(value='')
        self.var_tesseract_config = tk.StringVar(value='')
        self.var_status = tk.StringVar(value='Pronto')
        self.var_current_file = tk.StringVar(value='-')
        self.var_current_page = tk.StringVar(value='-')
        self.var_total_progress_label = tk.StringVar(value='0/0')
        self.var_page_progress_label = tk.StringVar(value='0/0')
        self.create_ui()
        self.after(100, self.process_ui_queue)

    def create_ui(self):
        root = ttk.Frame(self, padding=12)
        root.pack(fill='both', expand=True)
        form = ttk.LabelFrame(root, text='Configuração', padding=12)
        form.pack(fill='x')
        self.add_path_row(form, 0, 'Pasta de entrada', self.var_input_dir, self.browse_input_dir)
        self.add_path_row(form, 1, 'chunks.json', self.var_output_json, self.browse_output_json)
        self.add_path_row(form, 2, 'report.json', self.var_report_json, self.browse_report_json)
        self.add_path_row(form, 3, 'Tesseract', self.var_tesseract_cmd, self.browse_tesseract)
        row4 = ttk.Frame(form)
        row4.grid(row=4, column=0, columnspan=3, sticky='ew', pady=(8, 0))
        row4.columnconfigure((1, 3, 5), weight=1)
        ttk.Label(row4, text='Chunk size').grid(row=0, column=0, sticky='w')
        ttk.Entry(row4, textvariable=self.var_chunk_size, width=12).grid(row=0, column=1, sticky='ew', padx=(8, 18))
        ttk.Label(row4, text='Overlap').grid(row=0, column=2, sticky='w')
        ttk.Entry(row4, textvariable=self.var_overlap, width=12).grid(row=0, column=3, sticky='ew', padx=(8, 18))
        ttk.Label(row4, text='OCR idioma').grid(row=0, column=4, sticky='w')
        ttk.Entry(row4, textvariable=self.var_ocr_lang, width=16).grid(row=0, column=5, sticky='ew', padx=(8, 0))
        row5 = ttk.Frame(form)
        row5.grid(row=5, column=0, columnspan=3, sticky='ew', pady=(8, 0))
        row5.columnconfigure((1, 3, 5), weight=1)
        ttk.Label(row5, text='OCR DPI').grid(row=0, column=0, sticky='w')
        ttk.Entry(row5, textvariable=self.var_ocr_dpi, width=12).grid(row=0, column=1, sticky='ew', padx=(8, 18))
        ttk.Label(row5, text='OCR min chars').grid(row=0, column=2, sticky='w')
        ttk.Entry(row5, textvariable=self.var_ocr_min_chars, width=12).grid(row=0, column=3, sticky='ew', padx=(8, 18))
        ttk.Label(row5, text='Tesseract config').grid(row=0, column=4, sticky='w')
        ttk.Entry(row5, textvariable=self.var_tesseract_config).grid(row=0, column=5, sticky='ew', padx=(8, 0))
        actions = ttk.Frame(root, padding=(0, 12, 0, 12))
        actions.pack(fill='x')
        self.btn_start = ttk.Button(actions, text='Processar', command=self.start_processing)
        self.btn_start.pack(side='left')
        self.btn_stop = ttk.Button(actions, text='Parar', command=self.stop_processing, state='disabled')
        self.btn_stop.pack(side='left', padx=(8, 0))
        status_box = ttk.LabelFrame(root, text='Status', padding=12)
        status_box.pack(fill='x')
        status_grid = ttk.Frame(status_box)
        status_grid.pack(fill='x')
        status_grid.columnconfigure(1, weight=1)
        ttk.Label(status_grid, text='Status').grid(row=0, column=0, sticky='w')
        ttk.Label(status_grid, textvariable=self.var_status).grid(row=0, column=1, sticky='w', padx=(12, 0))
        ttk.Label(status_grid, text='Arquivo').grid(row=1, column=0, sticky='nw', pady=(6, 0))
        ttk.Label(status_grid, textvariable=self.var_current_file, wraplength=700).grid(row=1, column=1, sticky='w', padx=(12, 0), pady=(6, 0))
        ttk.Label(status_grid, text='Página').grid(row=2, column=0, sticky='w', pady=(6, 0))
        ttk.Label(status_grid, textvariable=self.var_current_page).grid(row=2, column=1, sticky='w', padx=(12, 0), pady=(6, 0))
        progress_box = ttk.LabelFrame(root, text='Progresso', padding=12)
        progress_box.pack(fill='x', pady=(12, 0))
        ttk.Label(progress_box, text='Arquivos').pack(anchor='w')
        self.total_progress = ttk.Progressbar(progress_box, orient='horizontal', mode='determinate')
        self.total_progress.pack(fill='x', pady=(4, 0))
        ttk.Label(progress_box, textvariable=self.var_total_progress_label).pack(anchor='e', pady=(2, 8))
        ttk.Label(progress_box, text='Páginas do arquivo atual').pack(anchor='w')
        self.page_progress = ttk.Progressbar(progress_box, orient='horizontal', mode='determinate')
        self.page_progress.pack(fill='x', pady=(4, 0))
        ttk.Label(progress_box, textvariable=self.var_page_progress_label).pack(anchor='e', pady=(2, 0))
        log_box = ttk.LabelFrame(root, text='Log', padding=12)
        log_box.pack(fill='both', expand=True, pady=(12, 0))
        self.log_text = tk.Text(log_box, wrap='word', height=18)
        self.log_text.pack(side='left', fill='both', expand=True)
        scrollbar = ttk.Scrollbar(log_box, orient='vertical', command=self.log_text.yview)
        scrollbar.pack(side='right', fill='y')
        self.log_text.configure(yscrollcommand=scrollbar.set)

    def add_path_row(self, parent, row, label, variable, button_command):
        frame = ttk.Frame(parent)
        frame.grid(row=row, column=0, columnspan=3, sticky='ew', pady=(0 if row == 0 else 8, 0))
        frame.columnconfigure(1, weight=1)
        ttk.Label(frame, text=label, width=14).grid(row=0, column=0, sticky='w')
        ttk.Entry(frame, textvariable=variable).grid(row=0, column=1, sticky='ew', padx=(8, 8))
        ttk.Button(frame, text='Selecionar', command=button_command).grid(row=0, column=2)

    def browse_input_dir(self):
        path = filedialog.askdirectory(initialdir=self.var_input_dir.get() or str(Path.cwd()))
        if path:
            self.var_input_dir.set(path)

    def browse_output_json(self):
        path = filedialog.asksaveasfilename(defaultextension='.json', filetypes=[('JSON', '*.json')], initialfile=Path(self.var_output_json.get()).name, initialdir=str(Path(self.var_output_json.get()).parent) if self.var_output_json.get() else str(Path.cwd()))
        if path:
            self.var_output_json.set(path)

    def browse_report_json(self):
        path = filedialog.asksaveasfilename(defaultextension='.json', filetypes=[('JSON', '*.json')], initialfile=Path(self.var_report_json.get()).name, initialdir=str(Path(self.var_report_json.get()).parent) if self.var_report_json.get() else str(Path.cwd()))
        if path:
            self.var_report_json.set(path)

    def browse_tesseract(self):
        path = filedialog.askopenfilename(filetypes=[('Executável', '*.exe'), ('Todos', '*.*')], initialdir=str(Path(self.var_tesseract_cmd.get()).parent) if self.var_tesseract_cmd.get() else str(Path.cwd()))
        if path:
            self.var_tesseract_cmd.set(path)

    def validate_inputs(self):
        try:
            chunk_size = int(self.var_chunk_size.get())
            overlap = int(self.var_overlap.get())
            ocr_dpi = int(self.var_ocr_dpi.get())
            ocr_min_chars = int(self.var_ocr_min_chars.get())
        except ValueError:
            raise ValueError('Chunk size, overlap, OCR DPI e OCR min chars precisam ser números inteiros.')
        if chunk_size <= 0:
            raise ValueError('Chunk size precisa ser maior que 0.')
        if overlap < 0:
            raise ValueError('Overlap não pode ser negativo.')
        if ocr_dpi <= 0:
            raise ValueError('OCR DPI precisa ser maior que 0.')
        if ocr_min_chars < 0:
            raise ValueError('OCR min chars não pode ser negativo.')
        return {
            'input_dir': self.var_input_dir.get().strip(),
            'output_json': self.var_output_json.get().strip(),
            'report_json': self.var_report_json.get().strip(),
            'chunk_size': chunk_size,
            'overlap': overlap,
            'ocr_lang': self.var_ocr_lang.get().strip() or 'por+eng',
            'ocr_dpi': ocr_dpi,
            'ocr_min_chars': ocr_min_chars,
            'tesseract_cmd': self.var_tesseract_cmd.get().strip(),
            'tesseract_config': self.var_tesseract_config.get().strip(),
        }

    def start_processing(self):
        if self.worker and self.worker.is_alive():
            return
        try:
            settings = self.validate_inputs()
        except Exception as e:
            messagebox.showerror('Erro', str(e))
            return
        self.log_text.delete('1.0', 'end')
        self.var_status.set('Iniciando...')
        self.var_current_file.set('-')
        self.var_current_page.set('-')
        self.var_total_progress_label.set('0/0')
        self.var_page_progress_label.set('0/0')
        self.total_progress['value'] = 0
        self.total_progress['maximum'] = 1
        self.page_progress['value'] = 0
        self.page_progress['maximum'] = 1
        self.processor = Processor(self.ui_queue, settings)
        self.worker = threading.Thread(target=self.processor.process, daemon=True)
        self.worker.start()
        self.btn_start.configure(state='disabled')
        self.btn_stop.configure(state='normal')

    def stop_processing(self):
        if self.processor:
            self.processor.request_stop()
            self.append_log('Solicitação de parada enviada.')

    def process_ui_queue(self):
        try:
            while True:
                event = self.ui_queue.get_nowait()
                kind = event.get('kind')
                if kind == 'log':
                    self.append_log(event['text'])
                elif kind == 'status':
                    self.var_status.set(event['text'])
                elif kind == 'file':
                    self.var_current_file.set(event['text'])
                elif kind == 'page':
                    self.var_current_page.set(event['text'])
                elif kind == 'total_progress':
                    current = max(0, int(event['current']))
                    total = max(1, int(event['total']))
                    self.total_progress['maximum'] = total
                    self.total_progress['value'] = current
                    self.var_total_progress_label.set(f'{current}/{total}')
                elif kind == 'page_progress':
                    current = max(0, int(event['current']))
                    total = max(1, int(event['total']))
                    self.page_progress['maximum'] = total
                    self.page_progress['value'] = current
                    self.var_page_progress_label.set(f'{current}/{total}')
                elif kind == 'done':
                    self.btn_start.configure(state='normal')
                    self.btn_stop.configure(state='disabled')
                    self.append_log(f"Concluído. Arquivos encontrados: {event['total_files_found']} | Arquivos processados: {event['total_files_processed']} | Entradas de texto: {event['total_text_entries']} | Chunks: {event['total_chunks']} | Páginas com OCR: {event['ocr_pages']}")
                    self.append_log(f"Saída: {event['output']}")
                    self.append_log(f"Relatório: {event['report']}")
                    messagebox.showinfo('Concluído', f"Processamento concluído.\n\nArquivos encontrados: {event['total_files_found']}\nArquivos processados: {event['total_files_processed']}\nEntradas de texto: {event['total_text_entries']}\nChunks: {event['total_chunks']}\nPáginas com OCR: {event['ocr_pages']}\n\nSaída:\n{event['output']}\n\nRelatório:\n{event['report']}")
                elif kind == 'error':
                    self.btn_start.configure(state='normal')
                    self.btn_stop.configure(state='disabled')
                    self.var_status.set('Erro')
                    self.append_log(f"Erro: {event['text']}")
                    messagebox.showerror('Erro', event['text'])
        except queue.Empty:
            pass
        if self.worker and not self.worker.is_alive():
            self.btn_start.configure(state='normal')
            self.btn_stop.configure(state='disabled')
        self.after(100, self.process_ui_queue)

    def append_log(self, text):
        self.log_text.insert('end', text + '\n')
        self.log_text.see('end')


def cli_mode():
    parser = argparse.ArgumentParser()
    parser.add_argument('--nogui', action='store_true')
    parser.add_argument('--input', default='')
    parser.add_argument('--output', default='chunks.json')
    parser.add_argument('--report', default='report.json')
    parser.add_argument('--chunk-size', type=int, default=1200)
    parser.add_argument('--overlap', type=int, default=200)
    parser.add_argument('--ocr-lang', default='por+eng')
    parser.add_argument('--ocr-dpi', type=int, default=250)
    parser.add_argument('--ocr-min-chars', type=int, default=30)
    parser.add_argument('--tesseract-cmd', default='')
    parser.add_argument('--tesseract-config', default='')
    args = parser.parse_args()
    if not args.nogui:
        return False
    if not args.input:
        raise ValueError('No modo --nogui, informe --input.')
    q = queue.Queue()
    settings = {
        'input_dir': args.input,
        'output_json': args.output,
        'report_json': args.report,
        'chunk_size': args.chunk_size,
        'overlap': args.overlap,
        'ocr_lang': args.ocr_lang,
        'ocr_dpi': args.ocr_dpi,
        'ocr_min_chars': args.ocr_min_chars,
        'tesseract_cmd': args.tesseract_cmd,
        'tesseract_config': args.tesseract_config,
    }
    processor = Processor(q, settings)
    processor.process()
    while not q.empty():
        event = q.get()
        if event['kind'] == 'log':
            print(event['text'])
        elif event['kind'] == 'done':
            print(f"Concluído. Arquivos encontrados: {event['total_files_found']} | Arquivos processados: {event['total_files_processed']} | Entradas de texto: {event['total_text_entries']} | Chunks: {event['total_chunks']} | OCR: {event['ocr_pages']}")
        elif event['kind'] == 'error':
            print(f"Erro: {event['text']}")
    return True


if __name__ == '__main__':
    try:
        ran_cli = cli_mode()
        if not ran_cli:
            app = App()
            app.mainloop()
    except Exception as e:
        print(f'Erro: {e}')
